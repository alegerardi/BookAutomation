import openai
import re
from docx import Document
from docx.shared import Pt, RGBColor
openai.api_key = YOUR_KEY_HERE                                                                                                    
mytheme = "Book on Large Language Models Evolution"
language = "english"


def docxConversion(input_file,output_file):
    doc = Document()
    is_a_Chapter     = False
    is_a_subchapter  = False

    with open(input_file,'r',encoding='utf-8') as file:
        for line in file:
            line = line.strip()



            if line == "CHAPTERFLAG":
                is_a_Chapter = True
                continue
            elif line == "SUBCHAPTERFLAG":
                is_a_subchapter = True
                continue



            if is_a_Chapter == True:
                doc.add_page_break()
                title1 = doc.add_heading(level=1)
                run = title1.add_run(line)
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0,0,0)
                run.font.bold = True


                is_a_Chapter = False
            elif is_a_subchapter == True:
                title2 = doc.add_heading(level=2)
                run = title2.add_run(line)
                run.font.size = Pt(17)
                run.font.color.rgb = RGBColor(0,0,0)
                run.font.bold = True

                is_a_subchapter = False
            else:
                doc.add_paragraph(line)
    doc.save(output_file)

def gerar_resposta100tokens(pergunta): 
    
    resposta = openai.ChatCompletion.create(
        model="gpt-4.1-mini", 
        messages=[
            {"role": "system", "content": "You are a book writer who double checks every single information and write in accordance to the language informed pontuactions."},  # Optional system message
            {"role": "user", "content": pergunta}  # User's input
        ],
        max_tokens=20  
    )
    return resposta['choices'][0]['message']['content'].strip()

def gerar_respostasNTokens(pergunta, N):

    resposta = openai.ChatCompletion.create(
        model="gpt-4.1-mini", 
        messages=[
            {"role": "system", "content": "You are a book writer who double checks every single information and write in accordance to the language informed pontuactions."},  # Optional system message
            {"role": "user", "content": pergunta}  # User's input
        ],
        max_tokens=N  
    )
    return resposta['choices'][0]['message']['content'].strip()

def remove_leading_numbers(file_path):
     with open(file_path, 'r+',encoding='utf-8') as file:
          lines = file.readlines()
          file.seek(0)
          file.truncate()
          for line in lines:
            cleaned_line = re.sub(r'^\d+\.\s*', '', line)
            file.write(cleaned_line)

def create_double_spaced_docx(input_file, output_file):
    # Create a new Document
    doc = Document()

    # Open the text file to read
    with open(input_file, 'r', encoding='iso-8859-1') as file:
        for line in file:
            line = line.strip()  # Remove any leading/trailing whitespace

            # Add a paragraph with double spacing
            paragraph = doc.add_paragraph(line)

            # Set the paragraph format to double spacing
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 2.0  # Set line spacing to double (2.0)
            paragraph_format.space_before = Pt(0)  # Optional: Adjust space before each paragraph
            paragraph_format.space_after = Pt(0)   # Optional: Adjust space after each paragraph

    # Save the document
    doc.save(output_file)
    print(f"Document saved as {output_file}")



def safe_filename(name):
    return re.sub(r'[\\/*?:"<>|]', '-', name)

mytheme = safe_filename(mytheme)

with open('chapters.txt', 'w',encoding='utf-8') as file:                                                                           #GETTING MAIN CHAPTERS
    chapLineByLinePrompt = "//USE DIRECT RESPONSE STYLE// In a book, give me ten topics, directly, line by line, about the following theme: " + mytheme + " written in the language " + language
    chapLineByLine = gerar_respostasNTokens(chapLineByLinePrompt, 200)
    file.truncate(0)
    file.write(chapLineByLine)


goodChapters = False

while goodChapters == False:
    with open('chapters.txt', 'r',encoding='utf-8') as file:                                                                           #GETTING TOPICS PER CHAPTER
        with open('topics.txt','a',encoding='utf-8') as topicsFile:
            topicsFile.truncate(0)
            for i in range(10):
                tmpChapter = file.readline() 
              
                print("About to do 4 topics about chapter: " + tmpChapter)
                inputGpt = "//USE DIRECT RESPONSE STYLE//in a " + mytheme + " book, I divide the chapter: " + tmpChapter + ": in 4 subchapters. Write, line by line, 4 (FOUR) possible subchapters. Do not use special characters, only alphanumeric and '.'. Write in " + language + " language, and do NOT enumerate each subchapter"
                resposta = gerar_respostasNTokens(inputGpt,80)

                topicsFile.write(resposta + "\n")
                print("Just did 4 subchapters about "+tmpChapter + '\n')

        print("Do you want another topic selection? \n")
        char = input("Y for yes, N for no\n")
        if char == 'N':
            goodChapters = True


with open( 'topics.txt', 'r', encoding='utf-8') as file:                                                                           #PROCCESING TOPICS RESPONSE DATA
    with open('topicsProccessed.txt','w', encoding='utf-8') as outfile:
        lines = file.readlines()
        
        filtered_lines = [line for line in lines if re.match(r'^[a-zA-Z0-9]', line.lstrip())]
        outfile.writelines(filtered_lines)

print("Check correctness of topicsProcessed")
ans = False
while ans != True:
    char = input("Procced with creation? Y for yes, N for n\n")
    if char == 'Y': 
        ans = True
    else:
        exit(0)


with open('chapters.txt', 'r', encoding = 'utf-8') as chaptersFile:                                                                #ASSEMBLING FINAL BOOK
    with open('topicsProccessed.txt','r', encoding='utf-8') as topicsFile:
        with open('finalBookBARRIED.txt','w',encoding='utf-8') as bookFile:
            for i in range(10):                             #TOTAL OF 10 CHAPTERS
                tmpChapter2 = chaptersFile.readline()
                bookFile.write("CHAPTERFLAG\n")
                bookFile.write("Chapter " + tmpChapter2+'\n')
                for j in range(4):
                    subchap1 = topicsFile.readline()
                    bookFile.write("\nSUBCHAPTERFLAG")
                    bookFile.write("\n" + str(i+1) + "." + str(j+1) + ": " + subchap1 +"\n")
                 
                    inputGpt1 = "//USE DIRECT RESPONSE STYLE// In a " + mytheme + " book, in the chapter ="+tmpChapter2 +", write content for the part covering  theme =" + subchap1 + ". Write maximum of 1500 words, dont enumerate sentences. Add at most 5 newline characters. USE ASCII CHARACTERS ONLY .Dont mention the theme name, write content DIRECTLY, in the language " + language
                    oneSubchapt = gerar_respostasNTokens(inputGpt1,2000)
                    bookFile.write(oneSubchapt)
                    print("Just did chap num" + str(i+1)+"."+str(j+1) + " = " + subchap1)

                bookFile.write('\n\n\n\n')

with open('finalBookBARRIED.txt', 'r',encoding='utf-8') as preFile:                                                               #PROCCESSING FINAL BOOK 
        lines = preFile.readlines()
filtered_lns = [line for line in lines if not re.match(r'^[*#. ]', line)]
with open('finalBook.txt','w', encoding = 'utf-8') as file:
        file.writelines(filtered_lns)

docxConversion("finalBookBARRIED.txt", mytheme+".docx")


                    

